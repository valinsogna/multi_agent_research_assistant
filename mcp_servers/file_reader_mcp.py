"""
MCP Server per lettura e analisi documenti.

Supporta:
- PDF
- DOCX
- TXT/Markdown
- CSV (preview)

Segue le best practices MCP di Anthropic.
"""

import json
from pathlib import Path
from typing import Optional, List
from enum import Enum
from datetime import datetime

from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field, ConfigDict


# =============================================================================
# Server Initialization
# =============================================================================

mcp = FastMCP("file_reader_mcp")


# =============================================================================
# Pydantic Models
# =============================================================================

class ResponseFormat(str, Enum):
    """Formato output."""
    MARKDOWN = "markdown"
    JSON = "json"
    RAW = "raw"


class ReadFileInput(BaseModel):
    """Input per leggere un file."""
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    
    file_path: str = Field(
        ...,
        description="Percorso del file da leggere"
    )
    max_chars: int = Field(
        default=10000,
        description="Numero massimo di caratteri da restituire",
        ge=100,
        le=100000
    )
    response_format: ResponseFormat = Field(
        default=ResponseFormat.MARKDOWN,
        description="Formato output"
    )


class ReadPDFInput(BaseModel):
    """Input per leggere un PDF."""
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    
    file_path: str = Field(
        ...,
        description="Percorso del file PDF"
    )
    pages: Optional[str] = Field(
        default=None,
        description="Pagine da estrarre (es. '1-5', '1,3,5', None=tutte)"
    )
    max_chars: int = Field(
        default=15000,
        description="Numero massimo di caratteri",
        ge=100,
        le=100000
    )


class ListFilesInput(BaseModel):
    """Input per listare files in una directory."""
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    
    directory: str = Field(
        ...,
        description="Percorso della directory"
    )
    pattern: str = Field(
        default="*",
        description="Pattern glob (es. '*.pdf', '*.txt')"
    )
    recursive: bool = Field(
        default=False,
        description="Se True, cerca anche nelle sottocartelle"
    )


class ExtractTextInput(BaseModel):
    """Input per estrazione testo strutturato."""
    model_config = ConfigDict(str_strip_whitespace=True, extra="forbid")
    
    file_path: str = Field(
        ...,
        description="Percorso del file"
    )
    extract_sections: bool = Field(
        default=True,
        description="Se True, identifica e separa le sezioni"
    )
    include_metadata: bool = Field(
        default=True,
        description="Se True, include metadati del file"
    )


# =============================================================================
# Helper Functions
# =============================================================================

def get_file_metadata(path: Path) -> dict:
    """Estrae metadati di un file."""
    stat = path.stat()
    return {
        "name": path.name,
        "extension": path.suffix.lower(),
        "size_bytes": stat.st_size,
        "size_human": _human_size(stat.st_size),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat()
    }


def _human_size(size_bytes: int) -> str:
    """Converte bytes in formato leggibile."""
    for unit in ["B", "KB", "MB", "GB"]:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def parse_page_range(pages_str: str, max_pages: int) -> List[int]:
    """Parsa stringa di pagine in lista di indici."""
    result = set()
    
    for part in pages_str.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            start = int(start) if start else 1
            end = int(end) if end else max_pages
            result.update(range(start, min(end + 1, max_pages + 1)))
        else:
            page = int(part)
            if 1 <= page <= max_pages:
                result.add(page)
    
    return sorted(result)


def read_pdf(file_path: Path, pages: Optional[str] = None, max_chars: int = 15000) -> dict:
    """Legge un file PDF."""
    from pypdf import PdfReader
    
    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    
    # Determina quali pagine leggere
    if pages:
        page_numbers = parse_page_range(pages, total_pages)
    else:
        page_numbers = list(range(1, total_pages + 1))
    
    # Estrai testo
    text_parts = []
    chars_count = 0
    pages_read = []
    
    for page_num in page_numbers:
        if chars_count >= max_chars:
            break
            
        page = reader.pages[page_num - 1]  # 0-indexed
        page_text = page.extract_text() or ""
        
        if chars_count + len(page_text) > max_chars:
            remaining = max_chars - chars_count
            page_text = page_text[:remaining] + "\n[... troncato ...]"
        
        text_parts.append(f"--- Pagina {page_num} ---\n{page_text}")
        chars_count += len(page_text)
        pages_read.append(page_num)
    
    return {
        "total_pages": total_pages,
        "pages_read": pages_read,
        "chars_extracted": chars_count,
        "text": "\n\n".join(text_parts),
        "metadata": reader.metadata._data if reader.metadata else {}
    }


def read_docx(file_path: Path, max_chars: int = 10000) -> dict:
    """Legge un file DOCX."""
    from docx import Document
    
    doc = Document(str(file_path))
    
    paragraphs = []
    chars_count = 0
    
    for para in doc.paragraphs:
        if chars_count >= max_chars:
            paragraphs.append("[... troncato ...]")
            break
        
        text = para.text.strip()
        if text:
            paragraphs.append(text)
            chars_count += len(text)
    
    # Estrai anche tabelle
    tables_text = []
    for table in doc.tables:
        if chars_count >= max_chars:
            break
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
            chars_count += sum(len(c) for c in cells)
        tables_text.append("\n".join(rows))
    
    return {
        "paragraphs_count": len(doc.paragraphs),
        "tables_count": len(doc.tables),
        "text": "\n\n".join(paragraphs),
        "tables": tables_text,
        "properties": {
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "created": str(doc.core_properties.created) if doc.core_properties.created else None
        }
    }


def read_text_file(file_path: Path, max_chars: int = 10000) -> dict:
    """Legge un file di testo (txt, md, csv, etc.)."""
    try:
        content = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = file_path.read_text(encoding="latin-1")
    
    if len(content) > max_chars:
        content = content[:max_chars] + "\n\n[... troncato ...]"
    
    lines = content.split("\n")
    
    return {
        "total_lines": len(lines),
        "total_chars": len(content),
        "text": content
    }


# =============================================================================
# MCP Tools
# =============================================================================

@mcp.tool(
    name="file_reader_read",
    annotations={
        "title": "Read File",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False
    }
)
async def file_reader_read(params: ReadFileInput) -> str:
    """
    Legge il contenuto di un file (PDF, DOCX, TXT, MD, CSV).
    
    Rileva automaticamente il tipo di file dall'estensione e
    applica il parser appropriato.
    
    Args:
        params: ReadFileInput con file_path, max_chars, response_format
    
    Returns:
        Contenuto del file in formato markdown o JSON
    """
    path = Path(params.file_path).expanduser().resolve()
    
    if not path.exists():
        return f"## ❌ Errore\n\nFile non trovato: {params.file_path}"
    
    if not path.is_file():
        return f"## ❌ Errore\n\nIl percorso non è un file: {params.file_path}"
    
    ext = path.suffix.lower()
    metadata = get_file_metadata(path)
    
    try:
        if ext == ".pdf":
            result = read_pdf(path, max_chars=params.max_chars)
            content = result["text"]
        elif ext in [".docx", ".doc"]:
            result = read_docx(path, max_chars=params.max_chars)
            content = result["text"]
            if result["tables"]:
                content += "\n\n### Tabelle\n\n" + "\n\n---\n\n".join(result["tables"])
        else:
            result = read_text_file(path, max_chars=params.max_chars)
            content = result["text"]
        
        if params.response_format == ResponseFormat.JSON:
            return json.dumps({
                "file": str(path),
                "metadata": metadata,
                "content": content,
                "extraction_info": {k: v for k, v in result.items() if k != "text"}
            }, indent=2, ensure_ascii=False)
        
        elif params.response_format == ResponseFormat.RAW:
            return content
        
        else:  # MARKDOWN
            output = [
                f"## 📄 {metadata['name']}",
                f"*{metadata['size_human']} | {ext} | Modificato: {metadata['modified'][:10]}*",
                "",
                content
            ]
            return "\n".join(output)
            
    except Exception as e:
        return f"## ❌ Errore durante la lettura\n\n{str(e)}"


@mcp.tool(
    name="file_reader_read_pdf",
    annotations={
        "title": "Read PDF",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False
    }
)
async def file_reader_read_pdf(params: ReadPDFInput) -> str:
    """
    Legge un file PDF con controllo sulle pagine.
    
    Permette di specificare quali pagine estrarre usando
    range (1-5) o lista (1,3,5).
    
    Args:
        params: ReadPDFInput con file_path, pages, max_chars
    
    Returns:
        Testo estratto dal PDF con indicazione delle pagine
    """
    path = Path(params.file_path).expanduser().resolve()
    
    if not path.exists():
        return f"## ❌ Errore\n\nFile non trovato: {params.file_path}"
    
    if path.suffix.lower() != ".pdf":
        return f"## ❌ Errore\n\nIl file non è un PDF: {params.file_path}"
    
    try:
        result = read_pdf(path, pages=params.pages, max_chars=params.max_chars)
        
        output = [
            f"## 📄 {path.name}",
            f"*Totale pagine: {result['total_pages']} | Pagine lette: {len(result['pages_read'])}*",
            f"*Caratteri estratti: {result['chars_extracted']}*",
            "",
            result["text"]
        ]
        
        return "\n".join(output)
        
    except Exception as e:
        return f"## ❌ Errore durante la lettura del PDF\n\n{str(e)}"


@mcp.tool(
    name="file_reader_list",
    annotations={
        "title": "List Files",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False
    }
)
async def file_reader_list(params: ListFilesInput) -> str:
    """
    Lista i file in una directory.
    
    Supporta pattern glob per filtrare i risultati.
    
    Args:
        params: ListFilesInput con directory, pattern, recursive
    
    Returns:
        Lista di file con metadati
    """
    dir_path = Path(params.directory).expanduser().resolve()
    
    if not dir_path.exists():
        return f"## ❌ Errore\n\nDirectory non trovata: {params.directory}"
    
    if not dir_path.is_dir():
        return f"## ❌ Errore\n\nIl percorso non è una directory: {params.directory}"
    
    # Trova files
    if params.recursive:
        files = list(dir_path.rglob(params.pattern))
    else:
        files = list(dir_path.glob(params.pattern))
    
    # Solo files, non directories
    files = [f for f in files if f.is_file()]
    
    if not files:
        return f"## 📁 {dir_path.name}\n\nNessun file trovato con pattern: {params.pattern}"
    
    # Formatta output
    output = [
        f"## 📁 {dir_path.name}",
        f"*{len(files)} file trovati (pattern: {params.pattern})*",
        ""
    ]
    
    # Raggruppa per estensione
    by_ext = {}
    for f in sorted(files):
        ext = f.suffix.lower() or "(nessuna)"
        if ext not in by_ext:
            by_ext[ext] = []
        by_ext[ext].append(f)
    
    for ext, ext_files in sorted(by_ext.items()):
        output.append(f"### {ext} ({len(ext_files)})")
        for f in ext_files[:20]:  # Max 20 per tipo
            meta = get_file_metadata(f)
            relative = f.relative_to(dir_path) if params.recursive else f.name
            output.append(f"- `{relative}` ({meta['size_human']})")
        if len(ext_files) > 20:
            output.append(f"- ... e altri {len(ext_files) - 20}")
        output.append("")
    
    return "\n".join(output)


@mcp.tool(
    name="file_reader_extract_structured",
    annotations={
        "title": "Extract Structured Text",
        "readOnlyHint": True,
        "destructiveHint": False,
        "idempotentHint": True,
        "openWorldHint": False
    }
)
async def file_reader_extract_structured(params: ExtractTextInput) -> str:
    """
    Estrae testo strutturato da un documento, identificando sezioni.
    
    Utile per documenti con struttura (report, articoli, manuali).
    
    Args:
        params: ExtractTextInput con file_path, extract_sections, include_metadata
    
    Returns:
        Testo strutturato con sezioni identificate
    """
    path = Path(params.file_path).expanduser().resolve()
    
    if not path.exists():
        return f"## ❌ Errore\n\nFile non trovato: {params.file_path}"
    
    metadata = get_file_metadata(path) if params.include_metadata else {}
    
    try:
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            result = read_pdf(path, max_chars=20000)
            raw_text = result["text"]
        elif ext in [".docx", ".doc"]:
            result = read_docx(path, max_chars=20000)
            raw_text = result["text"]
        else:
            result = read_text_file(path, max_chars=20000)
            raw_text = result["text"]
        
        # Identifica sezioni (headers)
        if params.extract_sections:
            import re
            
            sections = []
            current_section = {"title": "Introduzione", "content": []}
            
            for line in raw_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                
                # Pattern per headers (numeri, maiuscole, markdown)
                is_header = (
                    re.match(r"^#{1,3}\s+", line) or  # Markdown
                    re.match(r"^[0-9]+(\.[0-9]+)*\s+[A-Z]", line) or  # 1.2 Title
                    (len(line) < 100 and line.isupper()) or  # ALL CAPS
                    re.match(r"^(Chapter|Section|Part)\s+", line, re.I)  # Keywords
                )
                
                if is_header:
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"title": line, "content": []}
                else:
                    current_section["content"].append(line)
            
            # Aggiungi ultima sezione
            if current_section["content"]:
                sections.append(current_section)
            
            # Formatta output
            output = []
            if params.include_metadata:
                output.append(f"## 📄 {metadata['name']}")
                output.append(f"*{metadata['size_human']} | {len(sections)} sezioni identificate*\n")
            
            for i, section in enumerate(sections, 1):
                output.append(f"### {i}. {section['title']}")
                content = " ".join(section["content"][:500])  # Max 500 parole
                output.append(f"{content}\n")
            
            return "\n".join(output)
        
        else:
            return raw_text
            
    except Exception as e:
        return f"## ❌ Errore\n\n{str(e)}"


# =============================================================================
# Main Entry Point
# =============================================================================

if __name__ == "__main__":
    mcp.run()
